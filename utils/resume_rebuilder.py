# utils/resume_rebuilder.py
# ─────────────────────────────────────────────────────────────────────────────
# Resume Rebuilder — Phase 2: Export edited resume as DOCX / PDF
#
# Takes the final edited resume_json dict and produces downloadable files.
#
# resume_json structure:
#   {
#       "name":           "Raghav Pahwa",         (optional)
#       "contact":        "email | phone | ...",   (optional)
#       "summary":        "...",
#       "experience":     "...",
#       "projects":       "...",
#       "skills":         "...",
#       "education":      "...",
#       "certifications": "...",                   (optional)
#   }
# ─────────────────────────────────────────────────────────────────────────────

from __future__ import annotations

import io
import os
import tempfile
from typing import Optional

# ── python-docx ───────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

# ── docx2pdf ─────────────────────────────────────────────────────────────────
try:
    from docx2pdf import convert as _docx2pdf_convert
    _PDF_VIA_DOCX2PDF = True
except ImportError:
    _PDF_VIA_DOCX2PDF = False

# ── reportlab fallback ────────────────────────────────────────────────────────
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    _REPORTLAB_AVAILABLE = True
except ImportError:
    _REPORTLAB_AVAILABLE = False


# ── Section ordering and display labels ───────────────────────────────────────
SECTION_ORDER = ["summary", "experience", "projects", "skills", "education", "certifications"]
SECTION_LABELS = {
    "summary":        "Professional Summary",
    "experience":     "Work Experience",
    "projects":       "Projects",
    "skills":         "Technical Skills",
    "education":      "Education",
    "certifications": "Certifications",
}


def _add_heading(doc: "Document", text: str, level: int = 1) -> None:
    """Add a styled heading paragraph to the document."""
    para = doc.add_heading(text, level=level)
    run = para.runs[0] if para.runs else para.add_run(text)
    run.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)  # deep indigo
    run.font.size = Pt(13 if level == 1 else 11)
    run.bold = True
    para.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    para.paragraph_format.space_after = Pt(2)


def _add_divider(doc: "Document") -> None:
    """Add a thin horizontal rule after a section heading."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run("─" * 80)
    run.font.color.rgb = RGBColor(0xC7, 0xD2, 0xFE)  # soft indigo
    run.font.size = Pt(6)


def _add_body_text(doc: "Document", text: str) -> None:
    """Add body text, preserving line breaks and bullet points."""
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        # Detect bullet lines
        is_bullet = stripped.startswith(("•", "-", "*", "▸", "▹", "►", ">"))
        if is_bullet:
            # Strip bullet char
            content = stripped.lstrip("•-*▸▹►> ").strip()
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(content)
        else:
            para = doc.add_paragraph()
            run = para.add_run(stripped)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.space_before = Pt(0)


def build_docx(resume_json: dict, candidate_name: str = "Resume") -> bytes:
    """
    Build a clean, professional DOCX from the edited resume_json.

    Args:
        resume_json:    Dict mapping section names to their text content.
        candidate_name: Candidate's name for the document title.

    Returns:
        Bytes of the generated .docx file.

    Raises:
        ImportError: if python-docx is not installed.
    """
    if not _DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required for DOCX export. "
            "Install it with: pip install python-docx"
        )

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Header: Candidate name ────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(candidate_name)
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.name = "Calibri"
    name_run.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)
    name_para.paragraph_format.space_after = Pt(4)

    # Contact line (if present)
    contact = resume_json.get("contact", "").strip()
    if contact:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(contact)
        contact_run.font.size = Pt(9)
        contact_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        contact_para.paragraph_format.space_after = Pt(8)

    # ── Sections ──────────────────────────────────────────────────────────────
    for key in SECTION_ORDER:
        text = resume_json.get(key, "").strip()
        if not text:
            continue
        label = SECTION_LABELS.get(key, key.title())
        _add_heading(doc, label.upper(), level=1)
        _add_divider(doc)
        _add_body_text(doc, text)
        # Small gap between sections
        gap = doc.add_paragraph()
        gap.paragraph_format.space_after = Pt(4)

    # ── Serialise to bytes ────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_pdf_from_docx(docx_bytes: bytes) -> Optional[bytes]:
    """
    Convert DOCX bytes → PDF bytes using docx2pdf (requires Microsoft Word).

    Returns:
        PDF bytes on success, or None if conversion failed.
    """
    if not _PDF_VIA_DOCX2PDF:
        return None

    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "resume.docx")
        pdf_path  = os.path.join(tmp, "resume.pdf")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        try:
            _docx2pdf_convert(docx_path, pdf_path)
            if os.path.exists(pdf_path):
                with open(pdf_path, "rb") as f:
                    return f.read()
        except Exception:
            pass
    return None


def build_pdf_reportlab(resume_json: dict, candidate_name: str = "Resume") -> Optional[bytes]:
    """
    Fallback: generate a PDF directly with reportlab (no Word required).

    Returns:
        PDF bytes on success, or None if reportlab unavailable.
    """
    if not _REPORTLAB_AVAILABLE:
        return None

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=1.5 * cm, bottomMargin=1.5 * cm,
        leftMargin=2 * cm, rightMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    story = []

    # Name
    name_style = ParagraphStyle(
        "Name", parent=styles["Title"],
        fontSize=20, textColor=colors.HexColor("#1E1B4B"),
        spaceAfter=4, alignment=1,  # center
    )
    story.append(Paragraph(candidate_name, name_style))

    contact = resume_json.get("contact", "").strip()
    if contact:
        contact_style = ParagraphStyle(
            "Contact", parent=styles["Normal"],
            fontSize=9, textColor=colors.HexColor("#64748B"),
            spaceAfter=12, alignment=1,
        )
        story.append(Paragraph(contact, contact_style))

    heading_style = ParagraphStyle(
        "SectionHead", parent=styles["Heading1"],
        fontSize=11, textColor=colors.HexColor("#1E1B4B"),
        spaceBefore=12, spaceAfter=2, bold=True,
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontSize=10, textColor=colors.HexColor("#1E293B"),
        spaceAfter=3, leading=14,
    )
    bullet_style = ParagraphStyle(
        "Bullet", parent=styles["Normal"],
        fontSize=10, textColor=colors.HexColor("#1E293B"),
        spaceAfter=2, leading=14, leftIndent=14, bulletIndent=4,
    )

    for key in SECTION_ORDER:
        text = resume_json.get(key, "").strip()
        if not text:
            continue
        label = SECTION_LABELS.get(key, key.title()).upper()
        story.append(Paragraph(label, heading_style))
        story.append(HRFlowable(width="100%", thickness=0.5,
                                color=colors.HexColor("#C7D2FE"), spaceAfter=4))
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            is_bullet = stripped.startswith(("•", "-", "*", "▸", "▹", "►", ">"))
            content = stripped.lstrip("•-*▸▹►> ").strip() if is_bullet else stripped
            if is_bullet:
                story.append(Paragraph(f"• {content}", bullet_style))
            else:
                story.append(Paragraph(content, body_style))
        story.append(Spacer(1, 6))

    doc.build(story)
    buf.seek(0)
    return buf.read()


def export_resume(
    resume_json: dict,
    candidate_name: str = "Resume",
    fmt: str = "docx",
) -> tuple[bool, Optional[bytes], str]:
    """
    High-level export function called from app.py.

    Args:
        resume_json:    Edited resume dict.
        candidate_name: Candidate name for title.
        fmt:            "docx" or "pdf".

    Returns:
        (success, file_bytes, mime_type)
    """
    try:
        docx_bytes = build_docx(resume_json, candidate_name)
    except ImportError as e:
        return False, None, str(e)
    except Exception as e:
        return False, None, f"DOCX build error: {e}"

    if fmt == "docx":
        return True, docx_bytes, (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # PDF: try docx2pdf first, then reportlab
    pdf_bytes = build_pdf_from_docx(docx_bytes)
    if pdf_bytes:
        return True, pdf_bytes, "application/pdf"

    pdf_bytes = build_pdf_reportlab(resume_json, candidate_name)
    if pdf_bytes:
        return True, pdf_bytes, "application/pdf"

    return False, None, (
        "PDF export requires either Microsoft Word (for docx2pdf) "
        "or reportlab. Install one of them."
    )
